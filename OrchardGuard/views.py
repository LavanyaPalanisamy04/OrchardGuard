import ast
import csv
import json
import re

import firebase_admin
import requests
from django.conf import settings
from django.contrib import messages
from django.contrib.auth import logout
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.http import HttpResponseBadRequest, HttpResponse
from django.http import JsonResponse
from django.shortcuts import redirect
from django.shortcuts import render
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from firebase_admin import auth, credentials, db
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from django.contrib.auth.decorators import login_required


from .forms import FeedbackForm
from .forms import ListSearchForm, AnySearchForm
from .forms import SearchForm

cred = credentials.Certificate('D:\\pycharmproject\\djangoProject\\OrchardGuard\\security_key.json')
default_app = firebase_admin.initialize_app(cred, {
    'databaseURL': 'https://apple-disease-detection-ab165-default-rtdb.firebaseio.com/'
})

AWS_ELASTICSEARCH_URL = settings.AWS_ELASTICSEARCH_URL
AWS_ELASTICSEARCH_USERNAME = settings.AWS_ELASTICSEARCH_USERNAME
AWS_ELASTICSEARCH_PASSWORD = settings.AWS_ELASTICSEARCH_PASSWORD
auth = (AWS_ELASTICSEARCH_USERNAME, AWS_ELASTICSEARCH_PASSWORD)

COLUMNS_TO_EXPORT = ['acno', 'accession', 'cultivar_name', 'origin_country', 'origin_province', 'origin_city',
                     'e_pedigree', 'e_genus', 'e_species', 'breeder']
COLUMN_HEADER = ['Acno', 'Accession', 'Cultivar Name', 'Country', 'Province', 'City', 'Pedigree', 'Genus', 'Species',
                 'breeder']

REPORT_HEADING = 'Apple Accessions Report'
EXPORT_FILENAME = "Accessions Search Report"


def index_documents_opensearch(request):
    with open('OrchardGuard/AppleAccessions.json', 'r') as file:
        json_data = json.load(file)

    # Prepare bulk indexing request
    bulk_request = ''
    index_name = 'accessions'

    for document in json_data:
        # Add metadata line
        bulk_request += json.dumps({"index": {"_index": index_name}}) + '\n'
        # Add JSON document
        bulk_request += json.dumps(document) + '\n'

    # Send bulk indexing request
    headers = {'Content-Type': 'application/json'}
    response = requests.post(AWS_ELASTICSEARCH_URL, auth=auth, headers=headers, data=bulk_request)

    # Check response status
    if response.status_code == 200:
        print("Documents indexed successfully!")
    else:
        print("Failed to index documents:", response.text)


def list_search(request):
    if request.method == 'POST':
        form = ListSearchForm(request.POST)
        if form.is_valid():

            acnos = form.cleaned_data['acnos']

            # Now you can process the acnos data as needed
            query = {
                "query": {
                    "terms": {
                        "acno": [int(acno.strip()) for acno in acnos.split(',')]
                    }
                }
            }
            response = requests.post(
                AWS_ELASTICSEARCH_URL,
                auth=(AWS_ELASTICSEARCH_USERNAME, AWS_ELASTICSEARCH_PASSWORD),
                json=query)

            # Extract search results
            items = []
            for record in response.json()['hits']['hits']:
                items.append(record['_source'])

            return render(request, 'OrchardGuard/search_results.html', {'items': items})
    else:
        form = ListSearchForm()
    return render(request, 'OrchardGuard/search.html', {'form': form})


def any_search(request):
    if request.method == 'POST':
        form = AnySearchForm(request.POST)
        if form.is_valid():

            input = form.cleaned_data['input']

            # Now you can process the acnos data as needed
            query = {
                "query": {
                    "query_string": {
                        "query": input
                    }
                }
            }
            response = requests.post(
                AWS_ELASTICSEARCH_URL,
                auth=(AWS_ELASTICSEARCH_USERNAME, AWS_ELASTICSEARCH_PASSWORD),
                json=query)

            # Extract search results
            items = []
            for record in response.json()['hits']['hits']:
                items.append(record['_source'])

            return render(request, 'OrchardGuard/search_results.html', {'items': items})
    else:
        form = AnySearchForm()
    return render(request, 'OrchardGuard/search.html', {'form': form})


def elastic_search(request):
    if request.method == 'POST':
        form = SearchForm(request.POST)
        if form.is_valid():
            # Construct Elasticsearch query based on form data
            query = {
                "query": {
                    "bool": {
                        "should": []
                    }
                }
            }

            # Add search terms from the form fields
            for field, value in form.cleaned_data.items():
                if value:
                    query["query"]["bool"]["should"].append({"match": {field: value}})

                    # Make a POST request to Elasticsearch
            response = requests.post(
                AWS_ELASTICSEARCH_URL,
                auth=(AWS_ELASTICSEARCH_USERNAME, AWS_ELASTICSEARCH_PASSWORD),
                json=query)

            # Extract search results
            items = []
            for record in response.json()['hits']['hits']:
                items.append(record['_source'])
            # Render the search results in the template
            return render(request, 'OrchardGuard/search_results.html', {'items': items})
    else:
        form = SearchForm()
    return render(request, 'OrchardGuard/search.html', {'form': form})


def wrap_text(text, char_limit):
    words = text.split()
    wrapped_text = ''
    line = ''
    for word in words:
        if len(line) + len(word) <= char_limit:
            line += word + ' '
        else:
            wrapped_text += line + '\n'
            line = word + ' '
    wrapped_text += line
    return wrapped_text.strip()


def export_pdf(items_list):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{EXPORT_FILENAME}.pdf"'

    doc = SimpleDocTemplate(response, pagesize=letter)
    elements = []  # List to hold elements to add to the document
    # Define the styles for the document
    styles = getSampleStyleSheet()
    heading_style = styles['Heading1']  # Use a predefined heading style
    heading_style.alignment = 1
    # Create the heading Paragraph and add it to the elements list
    heading_text = REPORT_HEADING
    heading = Paragraph(heading_text, heading_style)
    elements.append(heading)

    # Define a style for wrapped text
    style = ParagraphStyle(name='WrapStyle', fontSize=8)

    # Create the header row with the column names
    header_row = [Paragraph('<b>' + column + '</b>', style) for column in COLUMN_HEADER]

    # Filter the items_list to only include the columns you want
    data = [header_row] + [
        [Paragraph(str(item.get(column, '')), style) for column in COLUMNS_TO_EXPORT] for item in items_list
    ]

    # Define column widths - you'll need to adjust these values to fit your content
    column_widths = [
        0.56 * inch, 0.75 * inch, 1 * inch, 0.75 * inch, 0.75 * inch,
        1 * inch, 0.75 * inch, 0.75 * inch, 0.75 * inch, 0.75 * inch
    ]

    # Make sure the number of widths matches the number of columns
    assert len(column_widths) == len(COLUMNS_TO_EXPORT), "Column widths do not match number of columns"

    # Create the table with the data and column widths
    table = Table(data, colWidths=column_widths)

    # Add style to table, including borders
    table_style = TableStyle([
        # ... your existing style definitions ...
        ('BOX', (0, 0), (-1, -1), 1, colors.black),  # Outer border
        ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Inner grid
    ])
    table.setStyle(table_style)
    elements.append(table)

    # List of elements to build the document with
    elements = [heading, Spacer(1, 0.25 * inch), table]  # Spacer adds space after the heading
    # Add table to the PDF
    doc.build(elements)
    return response


def export_word(items_list):
    # Create a Word document
    doc = Document()

    # Add a heading
    heading = doc.add_heading(REPORT_HEADING, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Add a table to the Word document
    table = doc.add_table(rows=1, cols=len(COLUMNS_TO_EXPORT))

    # Populate the header row
    for idx, column in enumerate(COLUMNS_TO_EXPORT):
        cell = table.cell(0, idx)
        cell.text = str(column)

    # Apply style to the table (optional, you can define your own style)
    table.style = 'Table Grid'

    # Populate table rows with items
    for item in items_list:
        row_cells = table.add_row().cells
        for idx, column in enumerate(COLUMNS_TO_EXPORT):
            row_cells[idx].text = str(item.get(column, ''))

    # Prepare the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{EXPORT_FILENAME}.docx"'

    # Save the document to the HttpResponse
    doc.save(response)

    return response


def export_csv(request):
    if request.method == 'POST':
        items_json = request.POST.get('items')
        export_option = request.POST.get('exportOption')

        try:
            # Parse the JSON data
            items_list = ast.literal_eval(items_json)

            # Filter the items_list to only include the columns you want
            filtered_items_list = [{column: item.get(column, '') for column in COLUMNS_TO_EXPORT} for item in
                                   items_list]

            if export_option == 'csv':
                # Export as CSV with selected columns
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = f'attachment; filename="{EXPORT_FILENAME}.csv"'

                csv_writer = csv.DictWriter(response, fieldnames=COLUMNS_TO_EXPORT)
                csv_writer.writeheader()
                for item in filtered_items_list:
                    csv_writer.writerow(item)

                return response
            elif export_option == 'pdf':
                # Export as PDF with selected columns
                return export_pdf(filtered_items_list)
            elif export_option == 'word':
                # Export as PDF with selected columns
                return export_word(filtered_items_list)
            else:
                return HttpResponseBadRequest("Invalid export option.")
        except (ValueError, SyntaxError) as e:
            return HttpResponseBadRequest("Invalid JSON data.")
    else:
        return HttpResponseBadRequest("Invalid request method.")


def upload_and_predict(request):
    if request.method == 'POST':
        if 'image_file' not in request.FILES:
            return JsonResponse({'error': 'No image file provided.'}, status=400)

        image_file = request.FILES['image_file']
        temp_image = default_storage.save("temp_image", ContentFile(image_file.read()))
        temp_image_path = default_storage.path(temp_image)

        try:
            with open(temp_image_path, 'rb') as file:
                fastapi_url = 'http://localhost:8081/prediction'
                response = requests.post(fastapi_url, files={'file': file})

            # Ensure you delete the temp file after closing it
            default_storage.delete(temp_image)

            if response.status_code == 200:
                prediction = response.json()
                return JsonResponse({
                    'class': prediction.get('class', 'N/A'),
                    'confidence': prediction.get('confidence', 0)
                })
            else:
                return JsonResponse({'error': 'Failed to get prediction from FastAPI.'}, status=response.status_code)
        except Exception as e:
            # If an error occurs, delete the temp file
            default_storage.delete(temp_image)
            return JsonResponse({'error': str(e)}, status=500)
    else:
        return render(request, 'OrchardGuard/image_recognition.html')


def homepage(request):
    return render(request, 'OrchardGuard/homepage.html')

def nonuser(request):
    return render(request, 'OrchardGuard/nonuser.html')


def feedback(request):
    if request.method == 'POST':
        form = FeedbackForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('feedback_success')  # Redirect to a success page after submitting feedback
    else:
        form = FeedbackForm()
    return render(request, 'OrchardGuard/feedback.html', {'form': form})


def information_page(request):
    any_search_form = AnySearchForm()
    search_form = SearchForm()
    list_search_form = ListSearchForm()

    return render(request, 'OrchardGuard/infohub.html', {
        'any_search_form': any_search_form,
        'search_form': search_form,
        'list_search_form': list_search_form
    })


def login(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')

        # Correct Firebase REST API endpoint for signing in with email and password
        url = f"https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword?key={settings.FIREBASE_WEB_API_KEY}"

        # Data should be sent in the body of the POST request as JSON
        data = {
            'email': email,
            'password': password,
            'returnSecureToken': True,
        }

        response = requests.post(url, json=data)  # Use json=data to send the payload as JSON

        if response.status_code == 200:
            # Login was successful, you can now redirect or set session data
            return redirect('/homepage/')
        else:
            # Decode the response to get the error message
            error_message = response.json().get('error', {}).get('message', 'Unknown error')
            messages.error(request, f'Login failed. Reason: {error_message}')
            return render(request, 'OrchardGuard/login.html')

    return render(request, 'OrchardGuard/login.html')



def signup(request):
    if request.method == "POST":
        email = request.POST.get('email')
        password = request.POST.get('password')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        dob = request.POST.get('dob')

        # Basic password validation (enforce this client-side, or use Django's validators)
        if len(password) < 8 or not re.search("[a-z]", password) or not re.search("[A-Z]", password) or not re.search(
                "[0-9]", password) or not re.search("[\W_]", password):
            messages.error(request, "Password does not meet the security requirements.")
            return render(request, 'OrchardGuard/signup.html')

        try:
            user = auth.create_user(email=email, password=password)
            # After the user is created, save the additional information in Realtime Database
            user_ref = db.reference(f'users/{user.uid}')
            user_ref.set({
                'first_name': first_name,
                'last_name': last_name,
                'dob': dob
            })
            messages.success(request, "Account created successfully")
            return redirect('login')
        except Exception as e:
            messages.error(request, f"Failed to create account: {str(e)}")
            return render(request, 'OrchardGuard/signup.html')
    else:
        return render(request,'OrchardGuard/signup.html')



# def login_or_register(request):
#     login_form = LoginForm()
#     register_form = RegistrationForm()
#     if request.method == 'POST':
#         if 'login_submit' in request.POST:
#             login_form = LoginForm(request.POST)
#             if login_form.is_valid():
#                 # Perform login logic here
#                 return redirect('homepage')  # Redirect to homepage after successful login
#         elif 'register_submit' in request.POST:
#             register_form = RegistrationForm(request.POST)
#             if register_form.is_valid():
#                 # Perform registration logic here
#                 # For example:
#                 # first_name = register_form.cleaned_data['first_name']
#                 # last_name = register_form.cleaned_data['last_name']
#                 # email = register_form.cleaned_data['email']
#                 # password = register_form.cleaned_data['password']
#                 return redirect('login_or_register')  # Redirect to login page after successful registration
#     return render(request, 'OrchardGuard/vinsha_login.html', {'login_form': login_form, 'register_form': register_form})


def logout_view(request):
    logout(request)
    # Redirect to a success page.
    return redirect('nonuser')
